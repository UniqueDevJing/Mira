"""DOCX 解析: Heading 样式 → title block; 文本框/SmartArt 优雅降级"""
import hashlib
import re
from pathlib import Path

from docx import Document as DocxDocument

from engines.parsing.models import UIRDocument


class DocxParser:
    supported_types = (".docx",)

    def parse(self, file_path: str) -> UIRDocument:
        doc = DocxDocument(file_path)
        blocks = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = p.style.name if p.style else ""
            if style.startswith(("Heading", "标题")):
                blocks.append({
                    "type": "title", "bbox": [], "content": text, "page_num": 1,
                    "metadata": {"heading_level": self._heading_level(style)},
                })
            else:
                blocks.append({
                    "type": "paragraph", "bbox": [], "content": text, "page_num": 1,
                    "metadata": {},
                })
        tables = []
        for t in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
            if rows:
                tables.append({"page_num": 1, "bbox": [], "matrix": rows, "headers": rows[0]})
        # 文本框/SmartArt/嵌套表格: doc.paragraphs 不返回, 自动跳过 (graceful degradation)
        doc_id = hashlib.sha256(Path(file_path).read_bytes()).hexdigest()[:16]
        return UIRDocument(doc_id=doc_id, source={"type": "docx", "path": file_path},
                           pages=[{"page_num": 1, "blocks": blocks}], tables=tables)

    @staticmethod
    def _heading_level(style: str) -> int:
        m = re.search(r"(\d+)", style)
        return min(int(m.group(1)), 6) if m else 1
