"""DocxParser: Heading 样式 → title block, 表格提取, 非段落内容降级"""

from docx import Document

from engines.parsing.docx_parser import DocxParser


def _make_docx(tmp_path) -> str:
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("这是正文段落。")
    doc.add_heading("1.1 背景", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "列A"
    table.rows[0].cells[1].text = "列B"
    table.rows[1].cells[0].text = "v1"
    table.rows[1].cells[1].text = "v2"
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return str(p)


def test_heading_levels(tmp_path):
    uir = DocxParser().parse(_make_docx(tmp_path))
    titles = [b for b in uir.pages[0]["blocks"] if b["type"] == "title"]
    assert [t["content"] for t in titles] == ["第一章 概述", "1.1 背景"]
    assert [t["metadata"]["heading_level"] for t in titles] == [1, 2]


def test_paragraph_block(tmp_path):
    uir = DocxParser().parse(_make_docx(tmp_path))
    paras = [b for b in uir.pages[0]["blocks"] if b["type"] == "paragraph"]
    assert any("正文段落" in b["content"] for b in paras)


def test_tables_extracted(tmp_path):
    uir = DocxParser().parse(_make_docx(tmp_path))
    assert len(uir.tables) == 1
    assert uir.tables[0]["matrix"][0] == ["列A", "列B"]
